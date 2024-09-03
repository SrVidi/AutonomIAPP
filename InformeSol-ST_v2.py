import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.runnables import RunnableSequence
from langchain.prompts import load_prompt
from docx import Document
import io
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Function to read the content of the uploaded file
def read_uploaded_file(file):
    if file.name.endswith('.docx'):
        doc = Document(file)
        return ' '.join([paragraph.text for paragraph in doc.paragraphs])
    else:
        return file.read().decode("utf-8")

# Load prompts from files
prompt_template_WRITER = load_prompt("PromptsV2/Writer.yaml")
prompt_template_TREATMENT = load_prompt("PromptsV2/Treatment.yaml")
prompt_template_REVISOR = load_prompt("PromptsV2/Revisor.yaml")
prompt_template_FINAL = load_prompt("PromptsV2/Final.yaml")
prompt_template_TRANSLATOR = load_prompt("PromptsV2/Translator.yaml")

# Streamlit app
def main():
    st.title("INFORME ALTA MÉDICA, IA W-R-F-T")

    uploaded_file = st.file_uploader("Subir archivo", type=["txt", "docx"])

    if uploaded_file is not None:
        file_content = read_uploaded_file(uploaded_file)
        google_api_key = st.text_input("Introduzca la clave API de Google", type="password")

        col1, col2 = st.columns(2)
        with col1:
            process_button = st.button("Genera Informe en Español")
        with col2:
            process_and_translate_button = st.button("Genera Informe en Català")

        if process_button or process_and_translate_button:
            if google_api_key:
                with st.spinner("Generando el informe de alta..."):
                    gemini_pro = ChatGoogleGenerativeAI(model="gemini-1.5-pro", google_api_key=google_api_key, temperature=0.1)
                    gemini_flash = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=google_api_key, temperature=0.1)

                    chain_1 = RunnableSequence(prompt_template_WRITER, gemini_pro)
                    chain_2 = RunnableSequence(prompt_template_TREATMENT, gemini_pro)
                    chain_3 = RunnableSequence(prompt_template_REVISOR, gemini_flash)
                    chain_4 = RunnableSequence(prompt_template_FINAL, gemini_pro)
                    
                    report_original = chain_1.invoke({"DOCUMENTOS_CLINICOS": file_content})
                    report_treatment = chain_2.invoke({"DOCUMENTOS_CLINICOS": file_content})
                    corrections = chain_3.invoke({"DOCUMENTOS_CLINICOS": file_content, "INFORME_ORIGINAL": report_original})
                    report_final = chain_4.invoke({"INFORME_ORIGINAL": report_original, "CORRECCIONES": corrections})

                    if process_and_translate_button:
                        chain_translator = RunnableSequence(prompt_template_TRANSLATOR, gemini_pro)
                        report_translated = chain_translator.invoke({"REPORT_FINAL": report_final.content})
                        treatment_translated = chain_translator.invoke({"REPORT_FINAL": report_treatment.content})
                        final_content = (report_translated.content, treatment_translated.content)
                    else:
                        final_content = (report_final.content, report_treatment.content)

                    # Display the final report
                    st.markdown(final_content[0])  # Display the main report
                    
                    st.markdown(final_content[1])  # Display the treatment report

                    st.session_state.final_content = final_content
                    st.session_state.gemini_pro = gemini_pro

                    st.button("Descargar Informe (Word)", key="download_button")
            else:
                st.error("Introduzca su clave API de Google")

        if st.session_state.get("download_button", False) and st.session_state.get("gemini_pro") is not None:
            with st.spinner("Styling document for download..."):
                final_content = st.session_state.final_content
                doc = Document()
                
                for content in final_content:
                    for paragraph in content.split('\n'):
                        if paragraph.strip() == "---" or paragraph.strip() == "":
                            continue
                        if paragraph.startswith('# '):
                            heading = doc.add_heading(paragraph[2:], level=1)
                            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in heading.runs:
                                run.font.color.rgb = RGBColor(0, 0, 139)
                        elif paragraph.startswith('## '):
                            heading = doc.add_heading(paragraph[3:], level=2)
                            for run in heading.runs:
                                run.font.color.rgb = RGBColor(0, 0, 139)
                        elif paragraph.startswith('### '):
                            heading = doc.add_heading(paragraph[4:], level=3)
                            for run in heading.runs:
                                run.font.color.rgb = RGBColor(0, 0, 139)
                        elif paragraph.startswith('* '):
                            p = doc.add_paragraph(style='List Bullet')
                            add_formatted_text(p, paragraph[2:])
                        else:
                            p = doc.add_paragraph()
                            add_formatted_text(p, paragraph)
                
                bio = io.BytesIO()
                doc.save(bio)
                
                st.download_button(
                    label="Report Final (Word)",
                    data=bio.getvalue(),
                    file_name="Informe_Alta.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

def add_formatted_text(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        if i % 2 == 0:
            paragraph.add_run(part)
        else:
            paragraph.add_run(part).bold = True

if __name__ == "__main__":
    main()
