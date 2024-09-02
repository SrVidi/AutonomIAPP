import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.runnables import RunnableSequence
from langchain.prompts import load_prompt
from docx import Document
import io
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Function to read the content of the uploaded file
def read_uploaded_file(file):
    if file.name.endswith('.docx'):
        doc = Document(file)  # Pass the file object directly
        return ' '.join([paragraph.text for paragraph in doc.paragraphs])
    else:
        return file.read().decode("utf-8")  # For text files, read the content directly

# Load prompts from files
prompt_template_WRITER = load_prompt("Prompts/Writer.yaml")
prompt_template_REVISOR = load_prompt("Prompts/Revisor.yaml")
prompt_template_CORRECTOR = load_prompt("Prompts/Corrector.yaml")
prompt_template_FINAL = load_prompt("Prompts/Final.yaml")
prompt_template_TRANSLATOR = load_prompt("Prompts/Translator.yaml")
prompt_template_STYLER = load_prompt("Prompts/Styler.yaml")

# Streamlit app
def main():
    st.title("INFORME ALTA MÉDICA, IA W-T-R-F")

    # File uploader
    uploaded_file = st.file_uploader("Choose a file", type=["txt", "docx"])

    if uploaded_file is not None:
        # Read the content of the uploaded file
        file_content = read_uploaded_file(uploaded_file)

        # Google API Key input
        google_api_key = st.text_input("Enter your Google API Key", type="password")

        col1, col2 = st.columns(2)
        with col1:
            process_button = st.button("Process Document")
        with col2:
            process_and_translate_button = st.button("Process and Translate Document")

        if process_button or process_and_translate_button:
            if google_api_key:
                with st.spinner("Processing document..."):
                    # Initialize language models
                    gemini_pro = ChatGoogleGenerativeAI(model="gemini-1.5-pro", google_api_key=google_api_key, temperature="0.1")
                    gemini_flash = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=google_api_key, temperature="0.1")

                    # Create chains
                    chain_1 = RunnableSequence(prompt_template_WRITER, gemini_pro)
                    chain_2 = RunnableSequence(prompt_template_REVISOR, gemini_flash)
                    chain_3 = RunnableSequence(prompt_template_CORRECTOR, gemini_flash)
                    chain_4 = RunnableSequence(prompt_template_FINAL, gemini_pro)
                    
                    # Generate reports
                    report_original = chain_1.invoke({"DOCUMENTOS_CLINICOS": file_content})
                    report_revised = chain_2.invoke({"DOCUMENTOS_CLINICOS": file_content, "INFORME_ORIGINAL": report_original})
                    report_corrected = chain_3.invoke({"INFORME_ORIGINAL": report_original, "CORRECCIONES": report_revised})
                    report_final = chain_4.invoke({"REPORT_FINAL": report_corrected})

                    if process_and_translate_button:
                        # Add translation chain
                        chain_translator = RunnableSequence(prompt_template_TRANSLATOR, gemini_pro)
                        report_translated = chain_translator.invoke({"REPORT_FINAL": report_final.content})
                        final_content = report_translated.content
                    else:
                        final_content = report_final.content

                    # Display the final report with Markdown formatting
                    st.markdown(final_content)

                    # Store the final content and models in session state
                    st.session_state.final_content = final_content
                    st.session_state.gemini_pro = gemini_pro

                    # Create a download button for the final report
                    st.button("Download Styled Report as DOCX", key="download_button")

            else:
                st.error("Please enter your Google API Key to process the document.")

        # Check if the download button was clicked
        if st.session_state.get("download_button", False) and st.session_state.get("gemini_pro") is not None:
            with st.spinner("Styling document for download..."):
                # Apply styling
                chain_styler = RunnableSequence(prompt_template_STYLER, st.session_state.gemini_pro)
                report_styled = chain_styler.invoke({"REPORT_FINAL": st.session_state.final_content})
                styled_content = report_styled.content

                # Create DOCX with styled content
                doc = Document()
                for paragraph in styled_content.split('\n'):
                    if paragraph.startswith('# '):
                        heading = doc.add_heading(paragraph[2:], level=1)
                        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in heading.runs:
                            run.font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue
                    elif paragraph.startswith('## '):
                        heading = doc.add_heading(paragraph[3:], level=2)
                        for run in heading.runs:
                            run.font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue
                    elif paragraph.startswith('### '):
                        heading = doc.add_heading(paragraph[4:], level=3)
                        for run in heading.runs:
                            run.font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue
                    elif paragraph.startswith('* '):
                        p = doc.add_paragraph(style='List Bullet')
                        add_formatted_text(p, paragraph[2:])
                    else:
                        p = doc.add_paragraph()
                        add_formatted_text(p, paragraph)
                
                bio = io.BytesIO()
                doc.save(bio)
                
                st.download_button(
                    label="Click here to download",
                    data=bio.getvalue(),
                    file_name="styled_final_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

def add_formatted_text(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        if i % 2 == 0:
            paragraph.add_run(part)
        else:
            paragraph.add_run(part).bold = True

if __name__ == "__main__":
    main()
