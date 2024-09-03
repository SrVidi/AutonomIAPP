import gradio as gr
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.runnables import RunnableSequence
from langchain.prompts import load_prompt
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
import os

# Function to read the content of the uploaded file
def read_uploaded_file(file):
    if file.name.endswith('.docx'):
        doc = Document(file)
        return ' '.join([paragraph.text for paragraph in doc.paragraphs])
    else:
        return file.read().decode("utf-8")

# Load prompts from files
prompt_template_WRITER = load_prompt("PromptsV2/Writer.yaml")
prompt_template_TREATMENT = load_prompt("PromptsV2/Treatment.yaml")
prompt_template_REVISOR = load_prompt("PromptsV2/Revisor.yaml")
prompt_template_FINAL = load_prompt("PromptsV2/Final.yaml")
prompt_template_TRANSLATOR = load_prompt("PromptsV2/Translator.yaml")

def generate_report(file, google_api_key, language):
    file_content = read_uploaded_file(file)
    
    gemini_pro = ChatGoogleGenerativeAI(model="gemini-1.5-pro", google_api_key=google_api_key, temperature=0.1)
    gemini_flash = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=google_api_key, temperature=0.1)

    chain_1 = RunnableSequence(prompt_template_WRITER, gemini_pro)
    chain_2 = RunnableSequence(prompt_template_TREATMENT, gemini_pro)
    chain_3 = RunnableSequence(prompt_template_REVISOR, gemini_flash)
    chain_4 = RunnableSequence(prompt_template_FINAL, gemini_pro)
    
    report_original = chain_1.invoke({"DOCUMENTOS_CLINICOS": file_content})
    report_treatment = chain_2.invoke({"DOCUMENTOS_CLINICOS": file_content})
    corrections = chain_3.invoke({"DOCUMENTOS_CLINICOS": file_content, "INFORME_ORIGINAL": report_original})
    report_final = chain_4.invoke({"INFORME_ORIGINAL": report_original, "CORRECCIONES": corrections})

    if language == "Català":
        chain_translator = RunnableSequence(prompt_template_TRANSLATOR, gemini_pro)
        report_translated = chain_translator.invoke({"REPORT_FINAL": report_final.content})
        treatment_translated = chain_translator.invoke({"REPORT_FINAL": report_treatment.content})
        final_content = (report_translated.content, treatment_translated.content)
    else:
        final_content = (report_final.content, report_treatment.content)

    return final_content

def create_word_document(final_content):
    doc = Document()
    
    for content in final_content:
        for paragraph in content.split('\n'):
            if paragraph.strip() == "---" or paragraph.strip() == "":
                continue
            if paragraph.startswith('# '):
                heading = doc.add_heading(paragraph[2:], level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 139)
            elif paragraph.startswith('## '):
                heading = doc.add_heading(paragraph[3:], level=2)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 139)
            elif paragraph.startswith('### '):
                heading = doc.add_heading(paragraph[4:], level=3)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 139)
            elif paragraph.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, paragraph[2:])
            else:
                p = doc.add_paragraph()
                add_formatted_text(p, paragraph)
    
    # Save the document to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        doc.save(tmp_file.name)
    return tmp_file.name

def add_formatted_text(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        if i % 2 == 0:
            paragraph.add_run(part)
        else:
            paragraph.add_run(part).bold = True

def process_and_generate(file, google_api_key, language):
    if not google_api_key:
        return "Please enter your Google API key.", None, gr.update(visible=False)
    
    if not language:
        return "Please select a language (Español or Català).", None, gr.update(visible=False)
    
    try:
        final_content = generate_report(file, google_api_key, language)
        
        # Format the report as Markdown
        report_text = "# Informe de Alta Médica\n\n"
        for content in final_content:
            report_text += content.replace('# ', '## ') + "\n\n"
        
        word_doc_path = create_word_document(final_content)
        return report_text, word_doc_path, gr.update(visible=True)
    except Exception as e:
        return f"An error occurred: {str(e)}", None, gr.update(visible=False)

# Gradio interface
with gr.Blocks() as iface:
    gr.Markdown("# INFORME ALTA MÉDICA, IA W-R-F-T")
    gr.Markdown("Upload a clinical document to generate a medical discharge report.")
    
    with gr.Row():
        file_input = gr.File(label="Upload file (txt or docx)")
        api_key = gr.Textbox(label="Google API Key", type="password")
        language = gr.Radio(["Español", "Català"], label="Language", value="Español")
    
    generate_button = gr.Button("Generate Report")
    
    with gr.Column():
        output_report = gr.Markdown(label="Generated Report")
        word_doc_output = gr.File(label="Download Word Document", visible=False)

    generate_button.click(
        process_and_generate,
        inputs=[file_input, api_key, language],
        outputs=[output_report, word_doc_output, word_doc_output]
    )

if __name__ == "__main__":
    iface.launch()
